from pathlib import Path
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available")

try:
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError as e:
    logger.warning(f"OCR libraries not available: {e}")
    OCR_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

try:
    from app.services.ocr_processor import get_ocr_processor
    OCR_PROCESSOR_AVAILABLE = True
except ImportError:
    OCR_PROCESSOR_AVAILABLE = False
    logger.warning("OCR processor not available")

class Document:
    """Simple document class to replace LangChain Document"""
    def __init__(self, page_content: str, metadata: Dict[str, Any] = None):
        self.page_content = page_content
        self.metadata = metadata or {}

class PDFLoader:
    """Load PDF files using PyPDF2 with OCR fallback"""
    def __init__(self, file_path: str, enable_ocr: bool = True):
        self.file_path = file_path
        self.enable_ocr = enable_ocr and OCR_PROCESSOR_AVAILABLE

    def load(self) -> List[Document]:
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed")

        docs = []
        try:
            with open(self.file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                has_text = False
                for i, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        has_text = True
                        docs.append(Document(page_content=text, metadata={"page": i + 1, "source": "text_extraction"}))

            if not has_text and self.enable_ocr:
                logger.info(f"No text found in PDF, attempting OCR on {self.file_path}")
                try:
                    ocr_processor = get_ocr_processor()

                    if ocr_processor.is_scanned_pdf(self.file_path):
                        logger.info("PDF appears to be scanned, using OCR")

                        texts = ocr_processor.extract_text_from_pdf_images(self.file_path)

                        for i, text in enumerate(texts):
                            if text.strip():
                                docs.append(Document(
                                    page_content=text,
                                    metadata={"page": i + 1, "source": "ocr"}
                                ))

                        logger.info(f"OCR extracted text from {len(docs)} pages")

                except Exception as ocr_error:
                    logger.warning(f"OCR failed: {ocr_error}. Returning text-based extraction if available.")

        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise

        return docs

class OCRPDFLoader:
    """Load PDF files using OCR"""
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        if not OCR_AVAILABLE:
            raise ImportError("OCR libraries not installed")

        pages = convert_from_path(self.file_path)
        docs = []
        for i, img in enumerate(pages):
            text = pytesseract.image_to_string(img)
            docs.append(Document(page_content=text, metadata={"page": i + 1}))
        return docs

class TextLoader:
    """Load text files"""
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        with open(self.file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": self.file_path})]

class DocxLoader:
    """Load DOCX files"""
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed")

        doc = DocxDocument(self.file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": self.file_path})]

def load_single_document(file_path: str) -> List[Document]:
    suffix = Path(file_path).suffix.lower()

    if suffix == '.pdf':
        try:
            loader = PDFLoader(file_path)
            docs = loader.load()
            if all(not doc.page_content.strip() for doc in docs):
                raise ValueError("No text extracted by PDFLoader")
            return docs
        except Exception as e:
            logger.info(f"PDF parsing failed, attempting OCR fallback: {e}")
            try:
                return OCRPDFLoader(file_path).load()
            except Exception as ocr_error:
                logger.error(f"OCR fallback failed: {ocr_error}")
                error_doc = Document(
                    page_content=f"[ERROR] Could not extract text from PDF: {str(e)}",
                    metadata={"error": True}
                )
                return [error_doc]
    elif suffix in ['.txt', '.md']:
        loader = TextLoader(file_path)
        return loader.load()
    elif suffix in ['.doc', '.docx']:
        loader = DocxLoader(file_path)
        return loader.load()
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

